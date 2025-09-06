import json, os, logging, re
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain.schema import Document as LC_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from sentence_transformers import SentenceTransformer, util

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

MAX_CHUNK_SIZE = 400
CLASSIFIED_DIR = "./data/classified"
FAISS_DIR = "./data/faiss"
os.makedirs(CLASSIFIED_DIR, exist_ok=True)
os.makedirs(FAISS_DIR, exist_ok=True)

# -------------------------
# 文件保存/加载
# -------------------------
def save_json(file_name: str, data: dict):
    path = os.path.join(CLASSIFIED_DIR, f"{file_name}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    logger.info("Saved JSON to %s", path)

def load_json(file_name: str) -> dict | None:
    path = os.path.join(CLASSIFIED_DIR, f"{file_name}.json")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

# -------------------------
# FAISS 保存 & 加载
# -------------------------
def save_faiss(file_name: str, db: FAISS):
    save_path = os.path.join(FAISS_DIR, file_name)
    os.makedirs(save_path, exist_ok=True)
    db.save_local(save_path)
    logger.info("Saved FAISS db to %s", save_path)

def load_faiss(file_name: str, embeddings_model=None) -> FAISS | None:
    save_path = os.path.join(FAISS_DIR, file_name)
    if os.path.exists(save_path):
        if embeddings_model is None:
            embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        return FAISS.load_local(save_path, embeddings_model, allow_dangerous_deserialization=True)
    return None

# -------------------------
# 文档读取
# -------------------------
def read_docx_paragraphs(docx_path: str):
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Document split into %d paragraphs", len(paragraphs))
    return paragraphs

# -------------------------
# FAISS 构建
# -------------------------
def semantic_split(text: str, max_size=MAX_CHUNK_SIZE):
    sentences = re.split(r"[。,.]", text.replace("\n"," "))
    sentences = [s.strip() for s in sentences if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip(): 
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip(): 
        sub_chunks.append(current.strip())
    return [sc for sc in sub_chunks if len(sc) > 5]

def build_faiss(structured_resume: dict, embeddings_model=None):
    from langchain.schema import Document as LC_Document
    docs = []
    for cat in ["work_experience", "projects", "education", "other"]:
        for entry in structured_resume.get(cat, []):
            text = entry.get("description") if isinstance(entry, dict) else str(entry)
            if not text: continue
            for sc in semantic_split(text):
                meta = {"category": cat}
                docs.append(LC_Document(page_content=sc, metadata=meta))
    if embeddings_model is None:
        embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if not docs: return None
    db = FAISS.from_documents(docs, embeddings_model)
    logger.info("FAISS database built with %d chunks", len(docs))
    return db

# -------------------------
# 查询 & 回填
# -------------------------
def auto_fill_with_query(structured_resume: dict, query_results: dict):
    mapping = {"工作经历": "work_experience","项目": "projects","教育": "education","技能": "skills"}
    for query, results in query_results.items():
        category = mapping.get(query)
        if not category: continue
        if category == "skills":
            for r in results:
                for skill in r.replace("。","").split(","):
                    skill = skill.strip()
                    if skill and skill not in structured_resume["skills"]:
                        structured_resume["skills"].append(skill)
        else:
            entries = structured_resume.get(category, [])
            for i, r in enumerate(results):
                if i < len(entries):
                    entries[i]["description"] = r
                else:
                    new_entry = {"description": r, "title": None, "company": None, "start_date": None, "end_date": None, "degree": None, "school": None, "skills": []}
                    entries.append(new_entry)
            structured_resume[category] = entries
    return structured_resume

def query_dynamic_category(db, structured_resume, query: str, top_k=10):
    docs = db.similarity_search(query, k=top_k*5)
    target_category = query  # 简化为直接使用 query
    candidate_paras = [doc.page_content for doc in docs[:top_k]]
    return {"query": query, "results": candidate_paras}

# -------------------------
# 解析简历段落 -> 结构化 JSON（最简版）
# -------------------------
def parse_resume_to_structured(paragraphs: list):
    structured = {"name": None,"email": None,"phone": None,"education":[],"work_experience":[],"projects":[],"skills":[],"other":[]}
    for para in paragraphs:
        para = para.strip()
        if not para: continue
        # 简化分类逻辑
        if any(k in para.lower() for k in ["university","college","bachelor","master","phd","degree","学士","硕士","博士"]):
            structured["education"].append({"description": para})
        elif any(k in para.lower() for k in ["project","build","develop","create"]):
            structured["projects"].append({"description": para})
        elif any(k in para.lower() for k in ["work","intern","engineer","manager","developer","consultant","实习","工作","任职"]):
            structured["work_experience"].append({"description": para})
        elif len(para)<100:
            structured["skills"].extend([para])
        else:
            structured["other"].append({"description": para})
    return structured

# -------------------------
# 主流程
# -------------------------
if __name__=="__main__":
    file_name = "Resume(AI).docx"
    file_path = f"./downloads/{file_name}"

    structured_resume = load_json(file_name)
    if structured_resume is None:
        paragraphs = read_docx_paragraphs(file_path)
        structured_resume = parse_resume_to_structured(paragraphs)
        save_json(file_name, structured_resume)

    # 构建 FAISS
    db = load_faiss(file_name)
    if db is None:
        db = build_faiss(structured_resume)
        save_faiss(file_name, db)

    # 测试查询 + 回填
    queries = ["工作经历","项目","教育","技能"]
    query_results = {}
    for q in queries:
        result = query_dynamic_category(db, structured_resume, q)
        query_results[q] = result["results"]

    structured_resume = auto_fill_with_query(structured_resume, query_results)
    save_json(file_name, structured_resume)

    print("\n===== FINAL STRUCTURED RESUME =====")
    print(json.dumps(structured_resume, ensure_ascii=False, indent=2))
