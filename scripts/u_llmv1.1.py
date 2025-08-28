import json
import re
import logging
from dotenv import load_dotenv
from docx import Document as DocxDocument

from langchain.schema import Document as LC_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# 这里假设你用的是 Ollama API 接本地 LLaMA3.2
import ollama
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

# ---------------- 配置 ----------------
MAX_CHUNK_SIZE = 400  # WorkExperience 段落可再切小块
LLM_MODEL_NAME = "llama-3.2"  # 你的本地 LLaMA3.2 模型
SIMILARITY_K = 5
# --------------------------------------

def read_docx_paragraphs(docx_path: str):
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Document split into %d paragraphs", len(paragraphs))
    return paragraphs

def classify_paragraphs(paragraphs: list):
    """
    使用本地 LLaMA3.2 将段落分类为：
    WorkExperience / Project / Education / Skills / Other
    返回分类字典
    """
    classified = {"WorkExperience": [], "Project": [], "Education": [], "Skills": [], "Other": []}

    # 强化系统提示，增加示例
    system_prompt = """
你是简历分类助手。请将输入段落严格分类为：
- WorkExperience: 工作经历
- Project: 项目经历
- Education: 教育经历
- Skills: 技能经历
- Other: 其他

**严格输出单行 JSON，不要附加任何文字或换行**，例如：
{"category": "WorkExperience"}

示例：
段落: "Worked at XYZ Corp as Data Scientist from Jan 2023 to Present"
输出: {"category": "WorkExperience"}

段落: "Developed AI chatbot for customer service"
输出: {"category": "Project"}

段落: "Bachelor of Science in Computer Science, MIT"
输出: {"category": "Education"}
"""

    # 正则备选方案
    category_keywords = {
        "WorkExperience": ["worked", "manager", "engineer", "internship", "experience", "lead", "responsible"],
        "Project": ["built", "developed", "created", "implemented", "project"],
        "Education": ["bachelor", "master", "university", "college", "degree"],
        "Skills": ["python", "tensorflow", "pytorch", "sql", "machine learning", "skills"]
    }

    for para in paragraphs:
        para_clean = para.strip().replace("\n", " ")
        category = "Other"

        try:
            response = ollama.chat(
                model="llama3.2:3b",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": para_clean}
                ]
            )

            raw_content = response.message.content
            try:
                data = json.loads(raw_content)
                category = data.get("category", "Other")
                if category not in classified:
                    category = "Other"
            except json.JSONDecodeError:
                # 二次尝试：正则匹配关键字
                para_lower = para_clean.lower()
                for cat, keywords in category_keywords.items():
                    if any(k in para_lower for k in keywords):
                        category = cat
                        break
                logger.warning("JSON 解析失败，将段落归类为 %s: %s", category, para_clean[:50])

        except Exception as e:
            logger.error("分类段落时出错: %s", str(e))

        classified[category].append(para_clean)

    # 输出分类统计
    for cat, lst in classified.items():
        logger.info("%d paragraphs classified as %s", len(lst), cat)

    return classified

def semantic_split(text: str, max_size=MAX_CHUNK_SIZE):
    """
    按句号拆分段落，避免块过长
    """
    sentences = [s.strip() for s in text.replace("\n", " ").split("。") if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip():
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip():
        sub_chunks.append(current.strip())
    return [sc for sc in sub_chunks if len(sc) > 30]

def build_faiss(paragraphs: list, embeddings_model=None):
    """
    将 WorkExperience 或 Project 段落生成向量，存入 FAISS
    """
    docs = []
    for para in paragraphs:
        sub_chunks = semantic_split(para)
        for sc in sub_chunks:
            docs.append(LC_Document(page_content=sc))

    if embeddings_model is None:
        embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Using HuggingFaceEmbeddings for FAISS")

    db = FAISS.from_documents(docs, embeddings_model)
    logger.info("FAISS database built with %d chunks", len(docs))
    return db

if __name__ == "__main__":
    test_file = "Resume(AI).docx"
    paragraphs = read_docx_paragraphs(test_file)
    classified = classify_paragraphs(paragraphs)

    # 仅 WorkExperience + Project 建向量库
    work_paras = classified["WorkExperience"] + classified["Project"]
    db = build_faiss(work_paras)

    # 查询示例
    query = "work experience"
    results = db.similarity_search(query, k=len(work_paras))
    for i, res in enumerate(results, 1):
        logger.info("[RESULT %d] %s", i, res.page_content)
