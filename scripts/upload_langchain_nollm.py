import os
import logging
from storage_client import StorageClient
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LC_Document
from dotenv import load_dotenv
from docx import Document as DocxDocument

try:
    from langchain_openai import OpenAIEmbeddings
except ImportError:
    OpenAIEmbeddings = None

load_dotenv()

# ----------------- 日志配置 -----------------
logging.basicConfig(
    level=logging.INFO,  # 默认 INFO，可改成 DEBUG
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)
# --------------------------------------------

MAX_CHUNK_SIZE = 400  # 每个块最大字符数，可调

from docx import Document

def split_docx_by_visual(docx_path, min_font_size=12):
    doc = DocxDocument(docx_path)
    chunks = []
    current_chunk = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 获取段落的字体信息
        run = para.runs[0] if para.runs else None
        font_size = run.font.size.pt if run and run.font.size else 11  # 默认 11pt
        bold = run.font.bold if run else False
        italic = run.font.italic if run else False

        # 判断是否视觉标题
        is_visual_header = font_size >= min_font_size and bold

        # 空行或视觉标题开启新 chunk
        if is_visual_header and current_chunk:
            chunks.append("\n".join(current_chunk))
            current_chunk = []

        current_chunk.append(text)

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def semantic_split(chunk: str, max_size=MAX_CHUNK_SIZE):
    """
    按句号、换行等语义拆分，保证每块不超过 max_size
    """
    sentences = [s.strip() for s in chunk.replace("\n", " ").split("。") if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip():
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip():
        sub_chunks.append(current.strip())

    # 过滤掉过短的内容
    sub_chunks = [sc for sc in sub_chunks if len(sc) > 30]

    return sub_chunks


def upload_to_longchain(object_name: str):
    """
    上传文件到 LongChain，使用 docx 样式 + 语义分段
    """
    client = StorageClient()
    local_path = client.read_file(object_name)

    ext = os.path.splitext(local_path)[1].lower()

    # Step 1: 初步分块
    if ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(local_path)
        initial_chunks = [d.page_content for d in loader.load()]
    elif ext == ".docx":
        initial_chunks = split_docx_by_visual(local_path)
    elif ext == ".txt":
        with open(local_path, "r", encoding="utf-8") as f:
            text = f.read()
        initial_chunks = [line for line in text.splitlines() if line.strip()]
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    logger.info("Initial split into %d chunks", len(initial_chunks))

    # Step 2: 语义细分
    final_docs = []
    for idx, chunk in enumerate(initial_chunks, 1):
        logger.info("Processing initial chunk %d/%d (len=%d)", idx, len(initial_chunks), len(chunk))
        sub_chunks = semantic_split(chunk)
        if not sub_chunks:  # 如果没有细分，直接用整个 chunk
            sub_chunks = [chunk]
        for sc in sub_chunks:
            final_docs.append(LC_Document(page_content=sc))

    logger.info("Semantic split into %d chunks", len(final_docs))

    # Step 3: 选择 Embeddings
    openai_key = os.getenv("OPENAI_API_KEY")
    if openai_key and OpenAIEmbeddings is not None:
        embeddings = OpenAIEmbeddings(openai_api_key=openai_key)
        logger.info("Using OpenAIEmbeddings")
    else:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Using HuggingFaceEmbeddings")

    # Step 4: 构建向量数据库
    db = FAISS.from_documents(final_docs, embeddings)
    logger.info("File %s indexed in FAISS successfully", object_name)
    return db


if __name__ == "__main__":
    test_file = "Resume(AI).docx"
    db = upload_to_longchain(test_file)

    query = "work experience"
    results = db.similarity_search(query, k=5)

    for i, res in enumerate(results, 1):
        logger.info("[RESULT %d] %s", i, res.page_content)
