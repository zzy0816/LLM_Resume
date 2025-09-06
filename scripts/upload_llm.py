import json
import os
import logging
import re
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain.schema import Document as LC_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from difflib import SequenceMatcher
from sentence_transformers import SentenceTransformer, util

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

MAX_CHUNK_SIZE = 400
LLM_MODEL_NAME = "llama3.2:3b"

CLASSIFIED_DIR = "./data/classified"
FAISS_DIR = "./data/faiss"
os.makedirs(CLASSIFIED_DIR, exist_ok=True)
os.makedirs(FAISS_DIR, exist_ok=True)

# -------------------------
# 文件保存/加载
# -------------------------
def save_json(file_name: str, data: dict):
    path = os.path.join(CLASSIFIED_DIR, f"{file_name}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    logger.info("Saved JSON to %s", path)

def load_json(file_name: str) -> dict | None:
    path = os.path.join(CLASSIFIED_DIR, f"{file_name}.json")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None

# -------------------------
# FAISS 保存 & 加载
# -------------------------
def save_faiss(file_name: str, db: FAISS):
    save_path = os.path.join(FAISS_DIR, file_name)
    os.makedirs(save_path, exist_ok=True)
    db.save_local(save_path)
    logger.info("Saved FAISS db to %s", save_path)

def load_faiss(file_name: str, embeddings_model=None) -> FAISS | None:
    save_path = os.path.join(FAISS_DIR, file_name)
    if os.path.exists(save_path):
        if embeddings_model is None:
            embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        return FAISS.load_local(save_path, embeddings_model, allow_dangerous_deserialization=True)
    return None

# -------------------------
# 文档读取
# -------------------------
def read_docx_paragraphs(docx_path: str):
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Document split into %d paragraphs", len(paragraphs))
    return paragraphs

# -------------------------
# 全局语义模型
# -------------------------
semantic_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

def semantic_fallback(para_clean: str) -> str:
    """基于语义相似度的回退分类"""
    categories = {
        "work_experience": "工作经历，如实习、任职、团队领导、负责项目",
        "projects": "项目经历，如开发、实现、搭建、研究",
        "education": "教育经历，如学位、本科、硕士、大学",
        "skills": "技能，如Python、SQL、TensorFlow、PyTorch、机器学习",
        "other": "其他内容"
    }
    para_emb = semantic_model.encode(para_clean, convert_to_tensor=True)
    cat_embs = {cat: semantic_model.encode(desc, convert_to_tensor=True) for cat, desc in categories.items()}
    sims = {cat: float(util.cos_sim(para_emb, emb)) for cat, emb in cat_embs.items()}
    
    # 选 top-2 并考虑阈值差
    sorted_sims = sorted(sims.items(), key=lambda x: x[1], reverse=True)
    best_cat, best_score = sorted_sims[0]
    second_score = sorted_sims[1][1] if len(sorted_sims) > 1 else 0.0

    if best_score > 0.35 and (best_score - second_score) > 0.05:
        return best_cat
    return "other"


# -------------------------
# 正则兜底
# -------------------------
email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
phone_pattern = r"(\+?\d[\d\s\-\(\)]{7,20})"

def extract_basic_info(text: str) -> dict:
    result = {}
    email_match = re.search(email_pattern, text)
    phone_match = re.search(phone_pattern, text)
    if email_match:
        result["email"] = email_match.group()
    if phone_match:
        # 清理多余空格和括号
        phone_clean = re.sub(r"[\s\(\)]", "", phone_match.group())
        result["phone"] = phone_clean
    return result


# -------------------------
# 分类归一化
# -------------------------
def normalize_category(cat: str) -> str:
    mapping = {
        "work": "work_experience",
        "workexperience": "work_experience",
        "work_experience": "work_experience",
        "projects": "projects",
        "project": "projects",
        "education": "education",
        "edu": "education",
        "skills": "skills",
        "skill": "skills",
        "other": "other"
    }
    key = cat.lower().replace(" ", "").replace("_", "")
    return mapping.get(key, "other")

from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline

# -------------------------
# NER 模型加载（resume-ner）
# -------------------------
def load_ner_pipeline():
    tokenizer = AutoTokenizer.from_pretrained("yashpwr/resume-ner-bert-v2")
    model = AutoModelForTokenClassification.from_pretrained("yashpwr/resume-ner-bert-v2")
    ner_pipe = pipeline("ner", model=model, tokenizer=tokenizer, aggregation_strategy="simple")
    return ner_pipe

ner_pipeline = load_ner_pipeline()

# -------------------------
# 更稳健的技能规范化（替换原 normalize_skills）
# -------------------------
def normalize_skills(skills: list) -> list:
    skills_set = set()
    for s in skills:
        if not isinstance(s, str):
            continue
        s_clean = s.strip()
        if not s_clean:
            continue
        s_lower = s_clean.lower()
        if s_lower in ["sql","llm","aws","hugging","gpu","api"]:
            skills_set.add(s_lower.upper())
        else:
            # 保持首字母大写形式（但如果原本已全大写也保留）
            if s_clean.isupper():
                skills_set.add(s_clean)
            else:
                skills_set.add(s_clean.capitalize())
    return sorted(list(skills_set))

# -------------------------
# 文本归一化与相似度判断（用来去重）
# -------------------------
def normalize_text_for_compare(s: str) -> str:
    if not isinstance(s, str):
        return ""
    s = s.lower().strip()
    # 移除常见中英文标点并折叠空白
    for ch in "。。，,、：:；;.-–—()[]{}\"'`··\u3000":
        s = s.replace(ch, " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s

def similar_ratio(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()

def is_duplicate_para(new_para: str, existing_list: list, threshold: float = 0.85) -> bool:
    norm_new = normalize_text_for_compare(new_para)
    if not norm_new:
        return True
    for e in existing_list:
        # e 可能是 dict 或 str
        if isinstance(e, dict):
            text = e.get("description") or " ".join([str(v) for v in e.values()])
        else:
            text = str(e)
        norm_e = normalize_text_for_compare(text)
        if not norm_e:
            continue
        if norm_new in norm_e or norm_e in norm_new:
            return True
        if similar_ratio(norm_new, norm_e) >= threshold:
            return True
    return False

# -------------------------
# 技能从文本中提取（FAISS 返回的段落常常是 "Platforms & Tech: A, B"）
# -------------------------
def extract_skills_from_text(text: str) -> list:
    if not isinstance(text, str):
        return []
    # 优先按 ":" 分割（常见 "Platforms & Tech: Ollama, Hugging Face"）
    if ":" in text:
        after = text.split(":", 1)[1]
    else:
        after = text
    # 按常见分隔符切分
    parts = re.split(r"[,\|;/、;]+", after)
    skills = []
    for p in parts:
        p = p.strip()
        # 忽略太短或包含描述性关键词的项
        if len(p) <= 1:
            continue
        if any(k in p.lower() for k in ["platform", "tech", "languages", "tools", "frameworks"]):
            continue
        # 去掉末尾多余的句号/中文句号
        p = p.rstrip("。.")
        if p:
            skills.append(p)
    return skills

# -------------------------
# 分类段落
# -------------------------
CATEGORY_FIELDS = {
    "work_experience": ["title","company","start_date","end_date","description"],
    "education": ["school","degree","grad_date","description"],
    "projects": ["description"],
    "skills": ["skills"],
    "other": ["description"]
}

def classify_paragraphs(paragraph: str, structured: dict) -> tuple[str, dict]:
    para_clean = paragraph.strip().replace("\r", " ").replace("\n", " ")
    if not para_clean:
        return "other", {}

    para_lower = para_clean.lower()

    # ---- Step 1: 排除联系方式/基本信息 ----
    info = extract_basic_info(para_clean)
    if info:
        structured["email"] = structured.get("email") or info.get("email")
        structured["phone"] = structured.get("phone") or info.get("phone")
        return "basic_info", {}

    if any(k in para_lower for k in ["linkedin", "github", "电话", "邮箱"]):
        return "basic_info", {}

    # ---- Step 2: 判断类别并初始化 data ----
    def init_data_for_category(cat, text):
        fields = CATEGORY_FIELDS.get(cat, ["description"])
        d = {f: None for f in fields}
        if "description" in fields:
            d["description"] = text
        if "skills" in fields:
            d["skills"] = []
        return d

    category = None
    edu_keywords = ["university", "college", "学院", "大学", "bachelor", "master", "phd", "ma", "ms", "mba"]
    work_keywords = ["intern", "engineer", "manager", "responsible", "工作", "实习", "任职", "developer", "consultant"]

    if any(k in para_lower for k in edu_keywords):
        category = "education"
        data = init_data_for_category(category, para_clean)
        parts = [p.strip() for p in para_clean.split("|")]
        if len(parts) >= 2:
            data["school"] = parts[0]
            data["degree"] = parts[1]
        for p in parts:
            year_match = re.search(r"\b(19|20)\d{2}\b", p)
            if year_match:
                data["grad_date"] = data["grad_date"] or year_match.group()

    elif any(k in para_lower for k in work_keywords):
        category = "work_experience"
        data = init_data_for_category(category, para_clean)
        parts = [p.strip() for p in para_clean.split("|")]
        if len(parts) >= 2:
            data["company"] = parts[0]
            data["title"] = parts[1]
        for p in parts:
            date_match = re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|[0-9]{4})\s*–\s*(Present|[0-9]{4})", p, re.I)
            if date_match:
                data["start_date"] = date_match.group(1)
                data["end_date"] = date_match.group(2)
    else:
        category = None
        data = init_data_for_category("other", para_clean)

    # ---- Step 3: NER 辅助解析 ----
    try:
        ner_results = ner_pipeline(para_clean)
        for ent in ner_results:
            label = ent["entity_group"].lower()
            val = ent["word"].strip()
            if label == "per" and structured.get("name") is None:
                structured["name"] = val
            elif label == "org" and "company" in data and not data["company"]:
                data["company"] = val
            elif label == "title" and "title" in data and not data["title"]:
                data["title"] = val
            elif label == "edu" and "school" in data and not data["school"]:
                data["school"] = val
            elif label == "degree" and "degree" in data and not data["degree"]:
                data["degree"] = val
            elif label == "skill" and "skills" in data and val not in data["skills"]:
                data["skills"].append(val)
            elif label == "date":
                if "start_date" in data and not data["start_date"]:
                    data["start_date"] = val
                elif "end_date" in data and not data["end_date"]:
                    data["end_date"] = val
    except Exception as e:
        print(f"[NER ERROR] {e}")

    # ---- Step 4: 技能关键词补全 ----
    skill_keywords = [
        "python","sql","pandas","numpy","scikit","sklearn","tensorflow",
        "pytorch","keras","docker","kubernetes","aws","gcp","azure",
        "spark","hadoop","tableau","powerbi","llm","llama","hugging"
    ]
    if "skills" in data:
        for kw in skill_keywords:
            if kw in para_lower and kw not in data["skills"]:
                data["skills"].append(kw.upper() if kw in ["sql","llm","aws","hugging"] else kw.capitalize())

    # ---- Step 5: fallback ----
    if not category:
        category = semantic_fallback(para_clean)
        category = normalize_category(category)
        data = init_data_for_category(category, para_clean)

    return normalize_category(category), data

# -------------------------
# parse_resume_to_structured（修改版）
# -------------------------
def parse_resume_to_structured(paragraphs: list):
    structured = {
        "name": None,
        "email": None,
        "phone": None,
        "education": [],
        "work_experience": [],
        "projects": [],
        "skills": [],
        "other": []
    }

    for para in paragraphs:
        category, data = classify_paragraphs(para, structured)

        if category == "basic_info":
            continue
        elif category == "work_experience":
            structured["work_experience"].append(data)
        elif category == "projects":
            structured["projects"].append(data)
        elif category == "education":
            structured["education"].append(data)
        elif category == "skills":
            structured["skills"].extend(data.get("skills", []))
        else:
            structured["other"].append(data)

    # 技能去重 & 标准化
    structured["skills"] = normalize_skills(structured["skills"])

    # 自动补全字段
    structured = auto_fill_fields(structured)

    return structured

# -------------------------
# FAISS 分段（增强教育段落保留短文本）
# -------------------------
def semantic_split(text: str, max_size=MAX_CHUNK_SIZE):
    """
    将文本按句拆分为子块，用于 FAISS 插入
    - 对短段落（< max_size）保留
    - 支持中文句号和英文逗号分句
    """
    sentences = re.split(r"[。,.]", text.replace("\n"," "))
    sentences = [s.strip() for s in sentences if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip(): 
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip(): 
        sub_chunks.append(current.strip())
    # 保留短段落，避免教育信息丢失
    sub_chunks = [sc for sc in sub_chunks if len(sc) > 5]
    return sub_chunks

# -------------------------
# 自动补全工作/教育字段
# -------------------------
def auto_fill_fields(structured_resume: dict) -> dict:
    for cat, fields in CATEGORY_FIELDS.items():
        new_entries = []
        for entry in structured_resume.get(cat, []):
            # 如果 entry 是字符串，则包装为 dict
            if isinstance(entry, str):
                entry_dict = {f: None for f in fields}
                if "description" in fields:
                    entry_dict["description"] = entry
                if "skills" in fields:
                    entry_dict["skills"] = []
                entry = entry_dict
            # 如果 entry 是 dict，就确保包含所有字段
            if isinstance(entry, dict):
                for f in fields:
                    if f not in entry or entry[f] is None:
                        if f == "skills":
                            entry[f] = []
                        elif f == "start_date":
                            entry[f] = "Unknown"
                        elif f == "end_date":
                            entry[f] = "Present"
                        else:
                            entry[f] = "N/A"
                new_entries.append(entry)
            else:
                # 万一仍然是非法类型，转为 description dict
                entry_dict = {f: None for f in fields}
                if "description" in fields:
                    entry_dict["description"] = str(entry)
                for f in fields:
                    if entry_dict[f] is None:
                        if f == "skills":
                            entry_dict[f] = []
                        elif f == "start_date":
                            entry_dict[f] = "Unknown"
                        elif f == "end_date":
                            entry_dict[f] = "Present"
                        else:
                            entry_dict[f] = "N/A"
                new_entries.append(entry_dict)
        structured_resume[cat] = new_entries

    # 最后处理整体技能字段（保证为字符串列表并标准化）
    if "skills" in structured_resume:
        # 只保留字符串
        structured_resume["skills"] = [s for s in structured_resume["skills"] if isinstance(s, str)]
        structured_resume["skills"] = normalize_skills(structured_resume["skills"])

    return structured_resume

# -------------------------
# 使用 FAISS 结果补全并去重（新增 faiss_auto_fill）
# -------------------------
def faiss_auto_fill(structured_resume: dict, db, top_k: int = 10) -> dict:
    """
    使用 FAISS 补全：按类别查询、解析并追加到 structured_resume（含去重逻辑）
    """
    cat2query = {
        "work_experience": "工作经历",
        "projects": "项目经历",
        "education": "教育经历",
        "skills": "技能"
    }

    for cat, q in cat2query.items():
        res = query_dynamic_category(db, structured_resume, q, top_k=top_k)
        candidates = res.get("results", []) if res else []

        if not candidates:
            continue

        if cat == "skills":
            # 从每个候选段落提取技能 token，合并
            extracted = []
            for para in candidates:
                extracted.extend(extract_skills_from_text(para))
            # 合并到主技能列表，去重由 normalize_skills 处理
            structured_resume.setdefault("skills", [])
            # 只添加非重复新技能
            for sk in extracted:
                # 基本去重：按 normalize_text_for_compare 判断
                if not any(normalize_text_for_compare(sk) == normalize_text_for_compare(existing) for existing in structured_resume["skills"]):
                    structured_resume["skills"].append(sk)
            # 最终标准化
            structured_resume["skills"] = normalize_skills(structured_resume["skills"])
            continue

        # 对于 work_experience / projects / education：
        for para in candidates:
            # 规范化 para（去掉句末句号）
            para_clean = para.strip().rstrip("。.")
            # 跳过重复段落
            if is_duplicate_para(para_clean, structured_resume.get(cat, []), threshold=0.86):
                continue

            # 试解析工作经历常见格式 "Company | Title | Location | Start – End"
            new_entry = {f: None for f in CATEGORY_FIELDS.get(cat, ["description"])}
            new_entry["description"] = para_clean

            # 解析 common "|" 分隔格式
            parts = [p.strip().rstrip("。.") for p in para_clean.split("|") if p.strip()]
            if cat == "work_experience" and len(parts) >= 2:
                # company, title = parts[0], parts[1]
                new_entry["company"] = parts[0]
                # 有时 title 在 parts[1]，有时交换，尽量填 title 字段
                new_entry["title"] = parts[1]
                # 尝试从整段提取时间范围
                date_match = re.search(r'(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\b\d{4}\b)[\w\s\.]*?)\s*[–-]\s*(?P<end>Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\b\d{4}\b)[\w\s\.]*)', para_clean, re.I)
                if date_match:
                    new_entry["start_date"] = date_match.group("start").strip()
                    new_entry["end_date"] = date_match.group("end").strip()
                else:
                    new_entry["start_date"] = "Unknown"
                    new_entry["end_date"] = "Present"

            elif cat == "education" and len(parts) >= 1:
                # 尝试常见 "School | Degree | ... | Year" 格式
                new_entry["school"] = parts[0]
                if len(parts) >= 2:
                    new_entry["degree"] = parts[1]
                # 年份尝试抽取
                year_match = re.search(r"\b(19|20)\d{2}\b", para_clean)
                if year_match:
                    new_entry["grad_date"] = year_match.group()

            else:
                # projects / fallback：保持 description，start/end 留空由 auto_fill 填充
                pass

            structured_resume.setdefault(cat, [])
            structured_resume[cat].append(new_entry)

    # 最后再做字段补全与规范化
    structured_resume = auto_fill_fields(structured_resume)
    return structured_resume

# -------------------------
# FAISS 构建
# -------------------------
def build_faiss(structured_resume: dict, embeddings_model=None):
    docs = []
    for cat in ["work_experience", "projects", "education", "other"]:
        for entry in structured_resume.get(cat, []):
            fields_to_use = CATEGORY_FIELDS.get(cat, ["description"])
            text_fields = [str(entry.get(f, "")) for f in fields_to_use if entry.get(f)]
            if not text_fields:
                text_fields = [str(entry.get("description", "")) or str(entry)]
            text = " ".join(text_fields).strip()
            if not text:
                continue

            chunks = semantic_split(text)
            if not chunks:
                chunks = [text]

            for sc in chunks:
                meta_cat = normalize_category(cat)
                meta = {"category": meta_cat}
                docs.append(LC_Document(page_content=sc, metadata=meta))
                print(f"[FAISS INSERT] cat={meta_cat}, snippet={sc[:80]}")

    logger.info("Total docs to insert into FAISS: %d", len(docs))

    if embeddings_model is None:
        embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    if not docs:
        print("[FAISS WARN] no docs to insert into FAISS (docs list empty)")
        return None

    db = FAISS.from_documents(docs, embeddings_model)
    logger.info("FAISS database built with %d chunks", len(docs))
    return db

# -------------------------
# 查询接口
# -------------------------
def detect_query_category(query:str):
    query_lower = query.lower()
    if any(k in query_lower for k in ["work","experience","career","job","employment","工作经历"]):
        return "work_experience"
    elif any(k in query_lower for k in ["project","built","developed","created","项目"]):
        return "projects"
    elif any(k in query_lower for k in ["education","degree","university","school","bachelor","master","教育"]):
        return "education"
    elif any(k in query_lower for k in ["skill","skills","python","tensorflow","ml","pytorch","sql","技能"]):
        return "skills"
    else:
        return None
    
def query_dynamic_category(db, structured_resume, query: str, top_k=10, use_category_filter=True):
    """
    基于 FAISS 查询指定类别段落（严格类别过滤）
    """
    docs = db.similarity_search(query, k=top_k*5)
    print(f"[QUERY DEBUG] retrieved {len(docs)} docs for query='{query}'")

    candidate_paras = []
    target_category = normalize_category(detect_query_category(query))

    if use_category_filter and target_category:
        for doc in docs:
            doc_cat = normalize_category(doc.metadata.get("category", "other"))

            if target_category == "education":
                doc_lower = doc.page_content.lower()
                has_school = any(tok in doc_lower for tok in ["university","college","学院","大学"])
                has_degree = any(tok in doc_lower for tok in ["bachelor","master","phd","bs","ms","mba","学士","硕士","博士"])
                if doc_cat == "education" and (has_school or has_degree):
                    candidate_paras.append(doc.page_content)

            elif target_category == "skills":
                skill_keywords = [
                    "python","sql","pandas","numpy","scikit","sklearn","tensorflow",
                    "pytorch","keras","docker","kubernetes","aws","gcp","azure",
                    "spark","hadoop","tableau","powerbi","llm","llama","hugging"
                ]
                if doc_cat == "skills":
                    candidate_paras.append(doc.page_content)
                elif doc_cat == "other" and len(doc.page_content) < 150:
                    if any(k in doc.page_content.lower() for k in skill_keywords):
                        candidate_paras.append(doc.page_content)

            else:
                if doc_cat == target_category:
                    candidate_paras.append(doc.page_content)

            if len(candidate_paras) >= top_k:
                break

        if not candidate_paras:
            print(f"[QUERY DEBUG] No candidate paragraphs found for query='{query}' with strict category filter.")
            return {"query": query, "results": []}
    else:
        candidate_paras = [doc.page_content for doc in docs[:top_k]]

    for i, p in enumerate(candidate_paras):
        print(f"[QUERY RESULT] {i+1}. {p[:140]}")

    return {"query": query, "results": candidate_paras}

def fill_query_exact(structured: dict, query_results: dict) -> dict:
    """
    使用 query_results 完全覆盖原 JSON 对应类别，
    保留基础信息 (name/email/phone)
    """
    # 保留基础信息
    base_info = {k: structured.get(k) for k in ["name", "email", "phone"]}

    # 初始化空结构
    new_structured = {
        "name": base_info.get("name"),
        "email": base_info.get("email"),
        "phone": base_info.get("phone"),
        "education": [],
        "work_experience": [],
        "projects": [],
        "skills": [],
        "other": []
    }

    # ---- 工作经历 ----
    for para in query_results.get("工作经历", []):
        parts = [p.strip() for p in para.split("|")]
        entry = {
            "company": parts[0] if len(parts) > 0 else "N/A",
            "title": parts[1] if len(parts) > 1 else "N/A",
            "start_date": "Unknown",
            "end_date": "Unknown",
            "description": para
        }
        # 尝试解析时间范围
        date_match = re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})\s*[-–]\s*(Present|\d{4})", para, re.I)
        if date_match:
            entry["start_date"] = date_match.group(1)
            entry["end_date"] = date_match.group(2)
        new_structured["work_experience"].append(entry)

    # ---- 教育经历 ----
    for para in query_results.get("教育经历", []):
        parts = [p.strip() for p in para.split("|")]
        entry = {
            "school": parts[0] if len(parts) > 0 else "N/A",
            "degree": parts[1] if len(parts) > 1 else "N/A",
            "grad_date": "N/A",
            "description": para
        }
        # 尝试抽取年份
        year_match = re.search(r"\b(19|20)\d{2}\b", para)
        if year_match:
            entry["grad_date"] = year_match.group()
        new_structured["education"].append(entry)

    # ---- 项目经历 ----
    for para in query_results.get("项目经历", []):
        entry = {"description": para}
        new_structured["projects"].append(entry)

    # ---- 技能 ----
    for para in query_results.get("技能", []):
        extracted = extract_skills_from_text(para)
        new_structured["skills"].extend(extracted)

    # 技能去重 & 规范化
    new_structured["skills"] = normalize_skills(new_structured["skills"])

    # ---- 其他 ----
    for para in query_results.get("其他", []):
        new_structured["other"].append({"description": para})

    return new_structured

# -------------------------
# 主流程示例
# -------------------------
def main_pipeline(file_name: str, mode: str = "exact") -> dict:
    """
    主流程：解析简历 -> 构建FAISS -> 查询 -> 填充 -> 保存
    Args:
        file_name: 简历文件名
        mode: 填充模式，目前支持 "exact"
    Returns:
        structured_resume: 最终结构化JSON
    """
    file_path = f"./downloads/{file_name}"

    # 1️⃣ 加载或解析简历
    structured_resume = load_json(file_name)
    if structured_resume is None:
        paragraphs = read_docx_paragraphs(file_path)
        structured_resume = parse_resume_to_structured(paragraphs)
        structured_resume = auto_fill_fields(structured_resume)
        save_json(file_name, structured_resume)

    # 2️⃣ 构建或加载 FAISS
    db = load_faiss(file_name)
    if db is None:
        db = build_faiss(structured_resume)
        save_faiss(file_name, db)

    # 3️⃣ 查询 FAISS 并生成 query_results
    queries = ["工作经历", "项目经历", "教育经历", "技能"]
    query_results = {}
    for q in queries:
        res = query_dynamic_category(db, structured_resume, q, top_k=10)
        query_results[q] = res.get("results", [])

    # 4️⃣ 使用 query 结果填充结构化 JSON
    if mode == "exact":
        structured_resume = fill_query_exact(structured_resume, query_results)
    else:
        raise ValueError(f"Unsupported mode: {mode}")

    # 5️⃣ 保存最终 JSON
    save_json(file_name + "_faiss_confirmed", structured_resume)

    return structured_resume

if __name__ == "__main__":
    file_name = "Resume(AI).docx"
    result = main_pipeline(file_name, mode="exact")

    print("\n===== FINAL STRUCTURED RESUME JSON =====")
    print(json.dumps(result, ensure_ascii=False, indent=2))
