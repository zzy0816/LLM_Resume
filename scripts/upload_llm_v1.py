import json
import logging
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain.schema import Document as LC_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import ollama

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

MAX_CHUNK_SIZE = 400
LLM_MODEL_NAME = "llama3.1:8b"

# 读取 .docx 文件中的所有段落，并返回一个段落列表
def read_docx_paragraphs(docx_path: str):
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Document split into %d paragraphs", len(paragraphs))
    return paragraphs

# 调用 LLM（Ollama）+ 关键词回退，给每个段落打上分类标签
def classify_paragraphs(paragraphs: list):
    classified = {"WorkExperience": [], "Project": [], "Education": [], "Skills": [], "Other": []}

    system_prompt = """
你是简历分类助手。请将输入段落严格分类为：
- WorkExperience: 工作经历
- Project: 项目经历
- Education: 教育经历
- Skills: 技能经历
- Other: 其他

严格输出单行 JSON，例如：{"category": "WorkExperience"}
"""
    category_keywords = {
        "WorkExperience": ["worked", "manager", "engineer", "internship", "experience", "lead", "responsible"],
        "Project": ["built", "developed", "created", "implemented", "project"],
        "Education": ["bachelor", "master", "university", "college", "degree"],
        "Skills": ["python", "tensorflow", "pytorch", "sql", "machine learning", "skills"]
    }

    for idx, para in enumerate(paragraphs):
        para_clean = para.strip().replace("\n", " ")
        category = "Other"

        try:
            response = ollama.chat(
                model=LLM_MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": para_clean}
                ]
            )
            raw_content = response.message.content
            try:
                data = json.loads(raw_content)
                category = data.get("category", "Other")
                if category not in classified:
                    category = "Other"
            except json.JSONDecodeError:
                # 关键词回退
                para_lower = para_clean.lower()
                for cat, keywords in category_keywords.items():
                    if any(k in para_lower for k in keywords):
                        category = cat
                        break
                logger.warning("JSON解析失败，将段落归类为 %s: %s", category, para_clean[:50])
        except Exception as e:
            logger.error("分类段落时出错: %s", str(e))

        classified[category].append((idx, para_clean))

    for cat, lst in classified.items():
        logger.info("%d paragraphs classified as %s", len(lst), cat)

    return classified

# 把长文本切成较小的语义片段，保证向量检索时更精准。
def semantic_split(text: str, max_size=MAX_CHUNK_SIZE):
    sentences = [s.strip() for s in text.replace("\n", " ").split("。") if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip():
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip():
        sub_chunks.append(current.strip())
    return [sc for sc in sub_chunks if len(sc) > 30]

# 基于分类结果构建带类别标签的 FAISS 向量数据库
def build_faiss_with_category(classified_paragraphs, embeddings_model=None):
    docs = []
    for category, lst in classified_paragraphs.items():
        for idx, para in lst:
            sub_chunks = semantic_split(para)
            for sc in sub_chunks:
                docs.append(LC_Document(page_content=sc, metadata={"category": category}))

    if embeddings_model is None:
        embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Using HuggingFaceEmbeddings for FAISS")

    db = FAISS.from_documents(docs, embeddings_model)
    logger.info("FAISS database built with %d chunks", len(docs))
    return db

# 根据用户的查询文本，推测属于哪一类
def detect_query_category(query: str):
    query_lower = query.lower()
    if any(k in query_lower for k in ["work", "experience", "career", "job", "employment"]):
        return "WorkExperience"
    elif any(k in query_lower for k in ["project", "built", "developed", "created"]):
        return "Project"
    elif any(k in query_lower for k in ["education", "degree", "university", "school", "bachelor", "master"]):
        return "Education"
    elif any(k in query_lower for k in ["skill", "skills", "python", "tensorflow", "ml", "pytorch", "sql"]):
        return "Skills"
    else:
        return "Other"

# 结合 向量检索 + 目标类别过滤，返回最相关的段落，并调用 LLM 做摘要
def query_dynamic_category(db, query: str, top_k=10, use_category_filter=True):
    """
    查询向量数据库，并返回原文段落，可选择是否按类别过滤
    
    Args:
        db: FAISS 向量数据库
        query: 用户输入的查询文本
        top_k: 返回的段落数量
        use_category_filter: 是否根据 detect_query_category 过滤类别
        
    Returns:
        原文段落列表或格式化文本
    """
    docs = db.similarity_search(query, k=top_k * 3)  # 扩大检索范围

    if use_category_filter:
        target_category = detect_query_category(query)
        candidate_paras = [
            doc.page_content for doc in docs
            if doc.metadata.get("category") == target_category
        ][:top_k]

        if not candidate_paras:
            return f"No relevant content found in category {target_category}"
    else:
        # 不做类别过滤，直接取 top_k 最相似
        candidate_paras = [doc.page_content for doc in docs][:top_k]

    # 格式化输出为原文形式
    result_text = ""
    for i, para in enumerate(candidate_paras, 1):
        result_text += f"{i}. {para}\n\n"

    return result_text

# 直接返回某个类别的原始段落（不经检索，避免串类）
def get_category_paragraphs(classified_paragraphs, category: str):
    return [para for _, para in classified_paragraphs.get(category, [])]

# 输出某个类别下的所有原始段落，格式化编号，保证原文不改写
def summarize_full_category(classified_paragraphs, category: str):
    paras = [para for _, para in classified_paragraphs.get(category, [])]
    if not paras:
        return f"该简历中没有检测到 {category} 类内容。"
    result = f"=== {category} 原文内容 ===\n\n"
    for i, para in enumerate(paras, 1):
        result += f"{i}. {para}\n\n"
    return result

if __name__ == "__main__":
    test_file = r"D:\project\LLM_Resume\downloads\Resume(AI).docx"
    paragraphs = read_docx_paragraphs(test_file)
    classified = classify_paragraphs(paragraphs)

    # 使用全部段落生成 FAISS 数据库
    faiss_paras = [para for cat in classified.values() for idx, para in cat]
    db = build_faiss_with_category(classified)
    for query in ["work experience"]:
        print(f"\n=== 查询: {query} ===")
        result_text = query_dynamic_category(db, query, top_k=10, use_category_filter=True)
        print(result_text)
