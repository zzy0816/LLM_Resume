# upload_to_longchain.py
import os
from storage_client import StorageClient
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from dotenv import load_dotenv
from docx import Document as DocxDocument

try:
    from langchain_openai import OpenAIEmbeddings
except ImportError:
    OpenAIEmbeddings = None

load_dotenv()


# 标题列表，用于分块
HEADERS = [
    "Career Goal", "Education", "Skills", "Professional Experience",
    "Industry Experience", "Projects"
]


def split_docx_by_headers(docx_path: str):
    """
    按简历标题分块，每个标题下的内容作为一个 Document
    """
    doc = DocxDocument(docx_path)
    chunks = []
    current_chunk = []
    current_header = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 检查是否为标题
        is_header = any(text.startswith(h) for h in HEADERS)
        if is_header:
            if current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = []
            current_header = text
            current_chunk.append(text)
        else:
            current_chunk.append(text)
    if current_chunk:
        chunks.append("\n".join(current_chunk))
    return chunks


def upload_to_longchain(object_name: str):
    client = StorageClient()
    local_path = client.read_file(object_name)

    ext = os.path.splitext(local_path)[1].lower()

    if ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(local_path)
        documents = loader.load()
    elif ext == ".docx":
        paragraphs = split_docx_by_headers(local_path)
        documents = [Document(page_content=p) for p in paragraphs]
    elif ext == ".txt":
        with open(local_path, "r", encoding="utf-8") as f:
            text = f.read()
        paragraphs = [line for line in text.splitlines() if line.strip()]
        documents = [Document(page_content=p) for p in paragraphs]
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    print(f"[INFO] Split into {len(documents)} chunks")

    # Embeddings选择
    openai_key = os.getenv("OPENAI_API_KEY")
    if openai_key and OpenAIEmbeddings is not None:
        embeddings = OpenAIEmbeddings(openai_api_key=openai_key)
        print("[INFO] Using OpenAIEmbeddings")
    else:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        print("[INFO] Using HuggingFaceEmbeddings")

    db = FAISS.from_documents(documents, embeddings)
    print(f"[INFO] File {object_name} indexed in FAISS successfully")
    return db


if __name__ == "__main__":
    test_file = "Resume(AI).docx"
    db = upload_to_longchain(test_file)

    query = "work experience"
    results = db.similarity_search(query, k=5)

    for i, res in enumerate(results, 1):
        print(f"\n[RESULT {i}]")
        print(res.page_content)
