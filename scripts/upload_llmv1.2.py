import json
import logging
from dotenv import load_dotenv
from docx import Document as DocxDocument

from langchain.schema import Document as LC_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

import ollama

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

MAX_CHUNK_SIZE = 400
LLM_MODEL_NAME = "llama3.2:3b"
SIMILARITY_K = 5


def read_docx_paragraphs(docx_path: str):
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Document split into %d paragraphs", len(paragraphs))
    return paragraphs


def classify_paragraphs(paragraphs: list):
    classified = {"WorkExperience": [], "Project": [], "Education": [], "Skills": [], "Other": []}

    system_prompt = """
你是简历分类助手。请将输入段落严格分类为：
- WorkExperience: 工作经历
- Project: 项目经历
- Education: 教育经历
- Skills: 技能经历
- Other: 其他

严格输出单行 JSON，例如：{"category": "WorkExperience"}
"""
    category_keywords = {
        "WorkExperience": ["worked", "manager", "engineer", "internship", "experience", "lead", "responsible"],
        "Project": ["built", "developed", "created", "implemented", "project"],
        "Education": ["bachelor", "master", "university", "college", "degree"],
        "Skills": ["python", "tensorflow", "pytorch", "sql", "machine learning", "skills"]
    }

    for idx, para in enumerate(paragraphs):
        para_clean = para.strip().replace("\n", " ")
        category = "Other"

        try:
            response = ollama.chat(
                model=LLM_MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": para_clean}
                ]
            )
            raw_content = response.message.content
            try:
                data = json.loads(raw_content)
                category = data.get("category", "Other")
                if category not in classified:
                    category = "Other"
            except json.JSONDecodeError:
                para_lower = para_clean.lower()
                for cat, keywords in category_keywords.items():
                    if any(k in para_lower for k in keywords):
                        category = cat
                        break
                logger.warning("JSON解析失败，将段落归类为 %s: %s", category, para_clean[:50])
        except Exception as e:
            logger.error("分类段落时出错: %s", str(e))

        classified[category].append((idx, para_clean))

    for cat, lst in classified.items():
        logger.info("%d paragraphs classified as %s", len(lst), cat)

    return classified


def semantic_split(text: str, max_size=MAX_CHUNK_SIZE):
    sentences = [s.strip() for s in text.replace("\n", " ").split("。") if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip():
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip():
        sub_chunks.append(current.strip())
    return [sc for sc in sub_chunks if len(sc) > 30]


def build_faiss(paragraphs: list, embeddings_model=None):
    docs = []
    for para in paragraphs:
        sub_chunks = semantic_split(para)
        for sc in sub_chunks:
            docs.append(LC_Document(page_content=sc))

    if embeddings_model is None:
        embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Using HuggingFaceEmbeddings for FAISS")

    db = FAISS.from_documents(docs, embeddings_model)
    logger.info("FAISS database built with %d chunks", len(docs))
    return db


def query_with_llm(db, paragraphs, query: str, top_k=10):
    # 1️⃣ FAISS 检索 top_k
    docs = db.similarity_search(query, k=top_k)
    candidate_paras = [doc.page_content for doc in docs]

    # 2️⃣ LLM 聚合生成摘要
    llm_prompt = f"""
你是智能简历助手。
根据用户查询 "{query}"，请从下面的段落中：
1. 只保留与 query 高度相关的内容
2. 删除与 query 无关的 Education/Skills/Other 内容

段落如下：
{chr(10).join(candidate_paras)}

严格输出最终整理后的文本，不要返回原始段落或 JSON。
"""
    response = ollama.chat(
        model=LLM_MODEL_NAME,
        messages=[{"role": "user", "content": llm_prompt}]
    )
    return response.message.content


if __name__ == "__main__":
    test_file = r"D:\project\LLM_Resume\downloads\Resume(AI).docx"
    paragraphs = read_docx_paragraphs(test_file)
    classified = classify_paragraphs(paragraphs)

    # 使用全部段落生成 FAISS 数据库
    faiss_paras = [para for cat in classified.values() for idx, para in cat]
    db = build_faiss(faiss_paras)

    # 用户 query 示例
    query = "work experience"
    result_text = query_with_llm(db, faiss_paras, query, top_k=10)

    print("=== LLM 聚合结果 ===")
    print(result_text)
