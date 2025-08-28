import json
import logging
from dotenv import load_dotenv
from docx import Document as DocxDocument

from langchain.schema import Document as LC_Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

import ollama

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

MAX_CHUNK_SIZE = 400
LLM_MODEL_NAME = "llama3.1:8b"


def read_docx_paragraphs(docx_path: str):
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Document split into %d paragraphs", len(paragraphs))
    return paragraphs


def classify_paragraphs(paragraphs: list):
    classified = {"WorkExperience": [], "Project": [], "Education": [], "Skills": [], "Other": []}

    system_prompt = """
你是简历分类助手。请将输入段落严格分类为：
- WorkExperience: 工作经历
- Project: 项目经历
- Education: 教育经历
- Skills: 技能经历
- Other: 其他

严格输出单行 JSON，例如：{"category": "WorkExperience"}
"""
    category_keywords = {
        "WorkExperience": ["worked", "manager", "engineer", "internship", "experience", "lead", "responsible"],
        "Project": ["built", "developed", "created", "implemented", "project"],
        "Education": ["bachelor", "master", "university", "college", "degree"],
        "Skills": ["python", "tensorflow", "pytorch", "sql", "machine learning", "skills"]
    }

    for idx, para in enumerate(paragraphs):
        para_clean = para.strip().replace("\n", " ")
        category = "Other"

        try:
            response = ollama.chat(
                model=LLM_MODEL_NAME,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": para_clean}
                ]
            )
            raw_content = response.message.content
            try:
                data = json.loads(raw_content)
                category = data.get("category", "Other")
                if category not in classified:
                    category = "Other"
            except json.JSONDecodeError:
                # 关键词回退
                para_lower = para_clean.lower()
                for cat, keywords in category_keywords.items():
                    if any(k in para_lower for k in keywords):
                        category = cat
                        break
                logger.warning("JSON解析失败，将段落归类为 %s: %s", category, para_clean[:50])
        except Exception as e:
            logger.error("分类段落时出错: %s", str(e))

        classified[category].append((idx, para_clean))

    for cat, lst in classified.items():
        logger.info("%d paragraphs classified as %s", len(lst), cat)

    return classified


def semantic_split(text: str, max_size=MAX_CHUNK_SIZE):
    sentences = [s.strip() for s in text.replace("\n", " ").split("。") if s.strip()]
    sub_chunks = []
    current = ""
    for s in sentences:
        if len(current) + len(s) + 1 <= max_size:
            current += s + "。"
        else:
            if current.strip():
                sub_chunks.append(current.strip())
            current = s + "。"
    if current.strip():
        sub_chunks.append(current.strip())
    return [sc for sc in sub_chunks if len(sc) > 30]

def build_faiss_with_category(classified_paragraphs, embeddings_model=None):
    """
    构建带类别 metadata 的 FAISS
    """
    docs = []
    for category, lst in classified_paragraphs.items():
        for idx, para in lst:
            sub_chunks = semantic_split(para)
            for sc in sub_chunks:
                docs.append(LC_Document(page_content=sc, metadata={"category": category}))

    if embeddings_model is None:
        embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Using HuggingFaceEmbeddings for FAISS")

    db = FAISS.from_documents(docs, embeddings_model)
    logger.info("FAISS database built with %d chunks", len(docs))
    return db

def detect_query_category(query: str):
    query_lower = query.lower()
    if any(k in query_lower for k in ["work", "experience", "career", "job", "employment"]):
        return "WorkExperience"
    elif any(k in query_lower for k in ["project", "built", "developed", "created"]):
        return "Project"
    elif any(k in query_lower for k in ["education", "degree", "university", "school", "bachelor", "master"]):
        return "Education"
    elif any(k in query_lower for k in ["skill", "skills", "python", "tensorflow", "ml", "pytorch", "sql"]):
        return "Skills"
    else:
        return "Other"


def query_dynamic_category(db, query: str, top_k=10):
    # 动态确定类别
    target_category = detect_query_category(query)

    # FAISS 检索 top_k
    docs = db.similarity_search(query, k=top_k)

    # 筛选属于目标类别的子块
    candidate_paras = [doc.page_content for doc in docs if doc.metadata.get("category") == target_category]

    if not candidate_paras:
        return f"No relevant content found in category {target_category}"

    # LLM 聚合
    llm_prompt = f"""
你是智能简历助手。
请根据用户查询 "{query}" 只保留以下段落中与 {target_category} 相关的内容：
1. 删除所有与该类别无关的内容
2. 输出完整、可读摘要，不允许包含其他类别内容

段落如下：
{chr(10).join(candidate_paras)}
"""
    response = ollama.chat(
        model=LLM_MODEL_NAME,
        messages=[{"role": "user", "content": llm_prompt}]
    )
    return response.message.content

if __name__ == "__main__":
    test_file = r"D:\project\LLM_Resume\downloads\Resume(AI).docx"
    paragraphs = read_docx_paragraphs(test_file)
    classified = classify_paragraphs(paragraphs)

    # 使用全部段落生成 FAISS 数据库
    faiss_paras = [para for cat in classified.values() for idx, para in cat]
    db = build_faiss_with_category(classified)
    for query in ["work experience"]:
        print(f"\n=== 查询: {query} ===")
        result_text = query_dynamic_category(db, query, top_k=10)
        print(result_text)
